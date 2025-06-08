import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import io
import json
import asyncio
import os
import requests
import ssl # Added for NLTK download certificate handling
import nltk.data # Explicitly import nltk.data

# Load environment variables from .env file (for local development)
from dotenv import load_dotenv
load_dotenv()

# Streamlit page configuration MUST be the first Streamlit command in the script
st.set_page_config(page_title="AI Resume Optimizer", layout="centered")

# --- NLTK Resource Downloader for Deployment ---
# This function handles downloading NLTK data to a writable directory in the cloud
# environment if it's not already present.
# It also includes SSL context handling for certificate issues.
# Use st.cache_resource to ensure this function runs only once per deployment
@st.cache_resource
def download_nltk_resources():
    st.write("Initializing NLTK resource check and download...")

    try:
        # Create an unverified SSL context to bypass potential certificate issues during download
        _create_unverified_https_context = ssl._create_unverified_context
    except AttributeError:
        # Legacy Python that doesn't have _create_unverified_context
        pass
    else:
        # Handle target environment that doesn't have proper certificate verification
        ssl._create_default_https_context = _create_unverified_https_context

    # Define a path within the application's root directory for NLTK data
    # This is often writable in cloud deployment environments
    nltk_data_app_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'nltk_data')

    # Ensure the directory exists
    if not os.path.exists(nltk_data_app_path):
        os.makedirs(nltk_data_app_path)
        st.info(f"Created NLTK data directory at: {nltk_data_app_path}")

    # CRITICAL: Clear NLTK's default search paths and add only our designated path
    # This ensures NLTK looks exactly where we want it to in the cloud environment.
    nltk.data.path = [nltk_data_app_path] # Overwrite default paths

    # Set NLTK_DATA environment variable for the session (secondary measure)
    os.environ["NLTK_DATA"] = nltk_data_app_path
    st.info(f"Set NLTK_DATA environment variable to: {os.environ['NLTK_DATA']}")
    st.info(f"NLTK data paths set to: {nltk.data.path}")


    # Removed 'punkt' from required_nltk_data as we are now using regex tokenization
    required_nltk_data = ['stopwords', 'wordnet'] 
    for data_item in required_nltk_data:
        st.info(f"Attempting to download/verify NLTK resource '{data_item}'...")
        try:
            # Perform the download directly. nltk.download is smart enough to skip if already present.
            nltk.download(data_item, download_dir=nltk_data_app_path, quiet=True) 
            st.success(f"NLTK resource '{data_item}' downloaded/verified successfully.")
        except Exception as e: # Catch any error during the actual download process
            st.error(f"Error during NLTK download of '{data_item}': {e}. This resource is critical for NLP processing.")
            # If download fails, the app might not function correctly.
    
    st.write("NLTK resource check and download complete.")
    # Return a dummy value to indicate completion for st.cache_resource
    return True

# Call the NLTK downloader at the very beginning of the app
download_nltk_resources()


# --- Text Extraction Functions ---
def extract_text_from_pdf(pdf_file):
    """
    Extracts text from an uploaded PDF file.
    """
    text = ""
    try:
        pdf_reader = PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None
    return text

def extract_text_from_docx(docx_file):
    """
    Extracts text from an uploaded DOCX file.
    """
    text = ""
    try:
        document = Document(docx_file)
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return None
    return text

# --- NLP Preprocessing Function ---
def preprocess_text(text):
    """
    Cleans and preprocesses text for analysis.
    Steps: lowercase, remove non-alphanumeric, tokenize (using regex), remove stopwords, lemmatize.
    """
    if not text:
        return ""
    text = text.lower()
    # Removed nltk.word_tokenize and replaced with regex for robustness in deployment
    tokens = re.findall(r'\b\w+\b', text) # Use regex for word tokenization
    stop_words = set(stopwords.words('english')) # This is where 'stopwords' is used
    tokens = [word for word in tokens if word not in stop_words]
    lemmatizer = WordNetLemmatizer() # This is where 'wordnet' is used
    tokens = [lemmatizer.lemmatize(word) for word in tokens]
    return ' '.join(tokens)

# --- Core Optimization Functions ---
def get_document_similarity_and_vectorizer(resume_text, job_description_text):
    """
    Calculates the cosine similarity between the resume and job description
    after preprocessing, and returns the TF-IDF vectorizer.
    """
    preprocessed_resume = preprocess_text(resume_text)
    preprocessed_jd = preprocess_text(job_description_text)

    if not preprocessed_resume or not preprocessed_jd:
        return 0.0, None

    documents = [preprocessed_resume, preprocessed_jd]

    tfidf_vectorizer = TfidfVectorizer()
    tfidf_matrix = tfidf_vectorizer.fit_transform(documents)

    # Calculate cosine similarity
    similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    return similarity_score, tfidf_vectorizer

def get_keyword_suggestions(resume_text, job_description_text, tfidf_vectorizer):
    """
    Identifies keywords from the job description that are relevant
    but potentially missing or underrepresented in the resume.
    """
    if not tfidf_vectorizer:
        return []

    preprocessed_resume = preprocess_text(resume_text)
    preprocessed_jd = preprocess_text(job_description_text)

    feature_names = tfidf_vectorizer.get_feature_names_out()

    # Get TF-IDF scores for job description
    jd_vector = tfidf_vectorizer.transform([preprocessed_jd]).toarray()[0]
    # Filter out keywords that have a very low score in JD (less relevant)
    jd_keywords = {feature_names[i]: jd_vector[i] for i in jd_vector.argsort()[-100:][::-1] if jd_vector[i] > 0.01}

    # Get TF-IDF scores for resume
    resume_vector = tfidf_vectorizer.transform([preprocessed_resume]).toarray()[0]
    resume_keywords = {feature_names[i]: resume_vector[i] for i in resume_vector.argsort()[-100:][::-1] if resume_vector[i] > 0.01}

    suggestions = []
    for keyword, jd_score in jd_keywords.items():
        # Suggest if keyword is important in JD but not present in resume or significantly less important
        if keyword not in resume_keywords or resume_keywords[keyword] < jd_score * 0.5:
            suggestions.append(keyword)

    # Sort suggestions by their importance in the job description
    suggestions.sort(key=lambda x: jd_keywords.get(x, 0), reverse=True)

    return suggestions[:15] # Return top 15 most relevant suggestions

async def generate_optimized_resume_text(resume_content, job_description):
    """
    Uses the Gemini API to generate a fully optimized resume text in a structured format
    suitable for DOCX conversion.
    """
    st.info("Generating your fully optimized resume with AI (this may take a moment)...")

    prompt = f"""
    You are an expert resume writer and optimizer. Your task is to rewrite the provided resume to be highly optimized for the given job description.
    Focus on making the resume as relevant as possible, incorporating keywords naturally, using strong action verbs, and emphasizing quantifiable achievements where appropriate.

    **Instructions for the optimized resume content:**
    1.  **Structure:** Use clear section headers and bullet points. Each main section (e.g., "PROFESSIONAL SUMMARY", "EXPERIENCE", "SKILLS", "EDUCATION") should start with an all-caps header without any leading symbols or Markdown bolding.
    2.  **Sub-sections:** For roles within "EXPERIENCE" or "PROJECTS", use a format like "Job Title | Company Name | Dates" as a sub-header.
    3.  **Content:**
        * **Summary/Objective:** Provide a concise paragraph.
        * **Experience/Projects:** Use bullet points for achievements, starting with strong action verbs and including quantifiable results (e.g., "Increased X by Y%", "Reduced Z by A"). Integrate relevant keywords from the job description naturally within these bullets.
        * **Skills:** List technical skills clearly.
        * **Education:** List degrees, institutions, and dates.
    4.  **Tone:** Maintain a professional and impactful tone throughout.
    5.  **No Conversational Text:** The entire output should be the optimized resume content, ready to be parsed.

    ---
    **Original Resume:**
    {resume_content}

    ---
    **Job Description:**
    {job_description}

    ---
    **Optimized Resume Content (formatted for easy parsing):**
    """

    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": prompt}]
            }
        ]
    }

    # Retrieve API key from environment variables (either .env locally or Streamlit secrets)
    apiKey = os.getenv("GEMINI_API_KEY")

    if not apiKey:
        st.error("Gemini API key not found. Please ensure it's set as an environment variable (e.g., in a .env file locally, or as a Streamlit secret on deployment, named 'GEMINI_API_KEY').")
        return "API key missing."

    apiUrl = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={apiKey}"

    st.write("Attempting to call Gemini API for full resume generation...")
    try:
        response = requests.post(
            apiUrl,
            headers={'Content-Type': 'application/json'},
            data=json.dumps(payload),
            timeout=90 # Increased timeout to 90 seconds for longer generation tasks
        )
        st.write(f"API call initiated for full resume generation. Status Code: {response.status_code}")
        response.raise_for_status() # Raise an exception for HTTP errors (4xx or 5xx)

        st.write("API call successful. Attempting to parse JSON response...")
        result = response.json()
        st.write("JSON response parsed.")

        if result.get('candidates') and len(result['candidates']) > 0 and \
           result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts') and \
           len(result['candidates'][0]['content']['parts']) > 0:
            st.success("Optimized resume text generated!")
            return result['candidates'][0]['content']['parts'][0].get('text', '') # Ensure text is retrieved or empty string
        else:
            st.error("Could not generate optimized resume text. The AI response was empty or malformed.")
            st.json(result) # Display the raw result for debugging
            return "Could not generate optimized resume text. The AI response was empty or malformed. Check console for details."
    except requests.exceptions.Timeout:
        st.error("API call timed out (90 seconds). The server took too long to respond. This might be due to a slow internet connection or high load on the API. Please try again.")
        return "API call timed out."
    except requests.exceptions.ConnectionError as e:
        st.error(f"Network connection error: {e}. The app could not reach the Gemini API. Please check your internet connection, proxy settings, or firewall.")
        return f"Network connection error: {e}"
    except requests.exceptions.HTTPError as e:
        st.error(f"HTTP error occurred: {e.response.status_code} - {e.response.text}. This might indicate an issue with the API key or request. Response: {e.response.text[:200]}...")
        return f"HTTP error: {e.response.status_code}"
    except json.JSONDecodeError as e:
        st.error(f"Error decoding API response JSON: {e}. The API returned non-JSON or malformed data. Response text: {response.text}")
        return f"Error decoding API response: {e}"
    except Exception as e:
        st.error(f"An unexpected error occurred during AI resume generation: {e}")
        return f"An unexpected error occurred: {e}"

# --- DOCX Conversion Function ---
def convert_text_to_docx(text_content):
    """
    Converts plain text content (expected to have Markdown-like headings and lists)
    into a Word (.docx) document. This aims for an ATS-friendly, simple structure.
    """
    document = Document()

    # Set basic margins for better ATS readability
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Process the text content line by line to add to the document
    lines = text_content.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line: # Skip empty lines
            continue

        # Check for main section headings (e.g., PROFESSIONAL SUMMARY, EXPERIENCE)
        # These are expected to be all caps and without asterisks based on new prompt
        if stripped_line.isupper() and len(stripped_line.split()) <= 5 and not stripped_line.startswith(('**', '#', '-', '*')):
            document.add_heading(stripped_line, level=1)
        # Check for sub-headings like "Job Title | Company Name | Dates"
        elif '|' in stripped_line and stripped_line.count('|') >= 1 and not stripped_line.startswith(('*', '-')): # Ensure it's not a bullet mistakenly
            document.add_paragraph(stripped_line, style='Intense Quote') # Using 'Intense Quote' for a slightly different style for sub-headings
        # Check for bullet points (AI is prompted to use '-' or '*')
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            document.add_paragraph(stripped_line[2:].strip(), style='List Bullet')
        # Default to regular paragraph
        else:
            document.add_paragraph(stripped_line)

    # Save the document to a BytesIO object
    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io.getvalue()

# --- Streamlit UI ---
async def main_app():
    st.markdown(
        """
        <style>
            .main-header {
                font-size: 2.5em;
                color: #4CAF50;
                text-align: center;
                margin-bottom: 30px;
                font-family: 'Inter', sans-serif;
            }
            .centered-message {
                text-align: center;
                font-size: 1.2em;
                color: #555;
                margin-bottom: 20px;
            }
            .stFileUploader > div > div > button {
                background-color: #4CAF50;
                color: white;
                border-radius: 8px;
                padding: 10px 20px;
                font-size: 1.1em;
                transition: background-color 0.3s ease;
            }
            .stFileUploader > div > div > button:hover {
                background-color: #45a049;
            }
            .stTextInput label {
                font-size: 1.2em;
                color: #333;
            }
            .stButton > button {
                background-color: #007BFF;
                color: white;
                border-radius: 8px;
                padding: 10px 20px;
                font-size: 1.2em;
                margin-top: 20px;
                transition: background-color 0.3s ease;
            }
            .stButton > button:hover {
                background-color: #0056b3;
            }
            .result-section {
                background-color: #f0f8ff;
                border-left: 5px solid #007BFF;
                padding: 20px;
                border-radius: 10px;
                margin-top: 30px;
            }
            .similarity-score {
                font-size: 1.8em;
                font-weight: bold;
                color: #28a745;
                text-align: center;
                margin-bottom: 20px;
            }
            .suggestions-header {
                font-size: 1.5em;
                color: #333;
                margin-top: 25px;
                margin-bottom: 15px;
                border-bottom: 2px solid #eee;
                padding-bottom: 5px;
            }
            .stTextArea textarea {
                border-radius: 8px;
                border: 1px solid #ccc;
                padding: 10px;
                font-size: 1em;
            }
            body {
                font-family: 'Inter', sans-serif;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.markdown('<h1 class="main-header">📄 AI Resume Optimizer 🚀</h1>', unsafe_allow_html=True)
    st.markdown("<p class='centered-message'>Upload Resume, Paste the JD and BOOM - OPTIMIZE YOUR RESUME - it works.</p>", unsafe_allow_html=True)

    # File Uploader for Resume
    st.subheader("1. Upload Your Resume")
    uploaded_resume = st.file_uploader("Choose a PDF or DOCX file", type=["pdf", "docx"])

    resume_text = ""
    if uploaded_resume is not None:
        file_extension = uploaded_resume.name.split('.')[-1]
        if file_extension == "pdf":
            resume_text = extract_text_from_pdf(uploaded_resume)
        elif file_extension == "docx":
            resume_text = extract_text_from_docx(uploaded_resume)
        else:
            st.error("Unsupported file type. Please upload a PDF or DOCX file.")

        if resume_text:
            st.success("Resume uploaded and text extracted successfully!")
            with st.expander("View Extracted Resume Text (Snippet)"):
                st.text_area("Extracted Text", resume_text[:1000] + "..." if len(resume_text) > 1000 else resume_text, height=200, disabled=True)
        else:
            st.warning("Could not extract text from your resume. Please try another file or ensure it's not an image-based PDF.")

    # Text Area for Job Description
    st.subheader("2. Paste the Job Description")
    job_description = st.text_area("Paste the full job description here", height=300)

    # Optimize Button
    if st.button("Optimize My Resume!"):
        if not resume_text:
            st.warning("Please upload your resume first.")
        elif not job_description:
            st.warning("Please paste the job description first.")
        else:
            with st.spinner("Analyzing and optimizing your resume..."):
                # 1. Similarity Score
                similarity_score, tfidf_vectorizer = get_document_similarity_and_vectorizer(resume_text, job_description)

                st.markdown('<div class="result-section">', unsafe_allow_html=True)
                st.markdown(f'<p class="similarity-score">Resume-Job Description Alignment Score: {similarity_score * 100:.2f}%</p>', unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

                st.markdown('<div class="result-section">', unsafe_allow_html=True)
                st.markdown('<h3 class="suggestions-header">🎯 Keyword Suggestions</h3>', unsafe_allow_html=True)

                if similarity_score == 0.0 and (not preprocess_text(resume_text) or not preprocess_text(job_description)):
                    st.warning("Cannot calculate similarity or suggest keywords. Please ensure both documents contain sufficient relevant text.")
                else:
                    # 2. Keyword Suggestions
                    keyword_suggestions = get_keyword_suggestions(resume_text, job_description, tfidf_vectorizer)
                    if keyword_suggestions:
                        st.write("Consider adding or emphasizing these keywords from the job description in your resume:")
                        for i, keyword in enumerate(keyword_suggestions):
                            st.markdown(f"- **{keyword}**")
                    else:
                        st.info("Your resume already seems to cover many key terms from the job description!")
                st.markdown("</div>", unsafe_allow_html=True)

                # --- New Section for Optimized Resume ---
                st.markdown('<div class="result-section">', unsafe_allow_html=True)
                st.markdown('<h3 class="suggestions-header">✨ Your Optimized Resume Content ✨</h3>', unsafe_allow_html=True)
                # 3. Generate and Display Optimized Resume Text
                optimized_resume_content = await generate_optimized_resume_text(resume_text, job_description)
                st.markdown(optimized_resume_content) # Display as Markdown

                # Add download button for DOCX
                try:
                    docx_bytes = convert_text_to_docx(optimized_resume_content)
                    st.download_button(
                        label="Download Optimized Resume (Word/DOCX)",
                        data=docx_bytes,
                        file_name="optimized_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Download the AI-generated optimized resume content as a Word (.docx) file. This format is highly ATS-friendly."
                    )
                except Exception as e:
                    st.error(f"Error generating DOCX: {e}.")
                    st.info("Ensure the `python-docx` library is installed: `pip install python-docx`")

                st.info("Note: This tool provides optimized text content in a structured DOCX format. For highly customized visual layouts, you may still prefer to copy this content into a professional resume builder or your original template.")
    st.markdown("---")
    st.markdown("Developed with ❤️ for you by Sirram Charan.")

if __name__ == "__main__":
    asyncio.run(main_app())
